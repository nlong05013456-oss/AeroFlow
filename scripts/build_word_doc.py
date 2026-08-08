import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_report():
    doc = docx.Document()

    # Set page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Colors
    PRIMARY_COLOR = RGBColor(10, 37, 64)     # Deep Navy (#0A2540)
    SECONDARY_COLOR = RGBColor(0, 168, 204)  # Cyan (#00A8CC)
    DARK_TEXT = RGBColor(40, 40, 40)
    CODE_COLOR = RGBColor(180, 40, 40)

    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = SECONDARY_COLOR
        p.paragraph_format.space_after = Pt(20)
        return p

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        return p

    def add_body(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(10)
            r_bold.font.bold = True
            r_bold.font.color.rgb = DARK_TEXT
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.italic = italic
        run.font.color.rgb = DARK_TEXT
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(10)
            r_bold.font.bold = True
            r_bold.font.color.rgb = DARK_TEXT
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT
        return p

    def add_code_block(code_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = CODE_COLOR
        return p

    def add_box_note(title, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F0F4F8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r_t = p.add_run(f"📌 {title}\n")
        r_t.font.bold = True
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = PRIMARY_COLOR
        r_b = p.add_run(text)
        r_b.font.size = Pt(9.5)
        r_b.font.color.rgb = DARK_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- TITLE SECTION ---
    add_title("BÁO CÁO THUYẾT MINH THIẾT KẾ & PHÁT TRIỂN ĐỒ ÁN AEROFLOW")
    add_subtitle("Chuyên Sâu Toàn Bộ Mã Nguồn Module `src/` Thực Thi Luồng ELT Pipeline, Quy Trình Chuẩn Kỹ Thuật Dữ Liệu & Web Portal")

    # Meta Info Box
    tbl_meta = doc.add_table(rows=3, cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    meta_data = [
        [("Tên Dự Án:", " AeroFlow Aviation Analytics Platform"), ("Nền Tảng Đám Mây:", " Google Cloud Platform (GCP)")],
        [("Kiến Trúc Pipeline:", " Modern ELT (Extract - Load - Transform)"), ("Mã Nguồn Cốt Lõi:", " Thư mục `src/` (Python & BigQuery SQL)")],
        [("Quy Mô Dữ Liệu:", " > 3 Triệu Bản Ghi Chuyến Bay Real-time & Batch"), ("Ứng Dụng Web:", " Web Portal 2 Cổng (Customer & Admin)")]
    ]
    for row_idx, row in enumerate(meta_data):
        for col_idx, (label, val) in enumerate(row):
            cell = tbl_meta.cell(row_idx, col_idx)
            set_cell_background(cell, "F0F4F8")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(label)
            r1.font.bold = True
            r1.font.size = Pt(9.5)
            r1.font.color.rgb = PRIMARY_COLOR
            r2 = p.add_run(val)
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = DARK_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- SECTION 1 ---
    add_heading_1("1. PHÂN TÍCH BÀI TOÁN NGHIỆP VỤ & NGUYÊN TẮC THIẾT KẾ (PROBLEM DEFINITION)")
    add_body("Hàng không là một trong những ngành kinh tế dịch vụ phức tạp nhất thế giới, vận hành liên tục 24/7 với yêu cầu tuyệt đối về tính chính xác thời gian và an toàn. Sự cố trễ chuyến bay (Flight Delays) và hủy chuyến (Cancellations) tạo ra những tổn thất khổng lồ về mặt kinh tế và trải nghiệm khách hàng.")

    add_box_note("TƯ TƯỞNG CỐT LÕI CỦA QUY TRÌNH THIẾT KẾ DỮ LIỆU (DATA DESIGN METHODOLOGY)", 
                 "Theo chuẩn Data Engineering hiện đại: Bài toán Nghiệp vụ ➔ Bản vẽ ERD ➔ Mô hình Kho Dữ liệu Star Schema (DIM & FACT). Sau khi đã có 'Bản vẽ thiết kế khung kho', chúng ta mới tiến hành xác định Nguồn dữ liệu và Lập trình Cào nạp (Data Ingestion) để nạp đúng dữ liệu vào khung kho!")

    # --- SECTION 2 ---
    add_heading_1("2. BẢN VẼ KIẾN TRÚC & SƠ ĐỒ QUAN HỆ THỰC THỂ (CONCEPTUAL ERD)")
    add_body("Xác định 7 Thực thể chính: CHUYẾN BAY (Flight), SÂN BAY (Airport), HÃNG BAY (Carrier), TÀU BAY (Aircraft), THỜI TIẾT (Weather), THỜI GIAN (Date), LÝ DO HỦY (Cancellation Reason).")

    # --- SECTION 3 ---
    add_heading_1("3. THIẾT KẾ MÔ HÌNH KHO DỮ LIỆU CHUẨN (STAR SCHEMA: DIM & FACT)")
    add_body("Mô hình Star Schema gồm Bảng Sự kiện fact_flights (>2.95 Triệu dòng, Partition theo Ngày, Cluster theo Hãng/Sân bay) và 6 Bảng Chiều (dim_airport, dim_carrier, dim_aircraft, dim_weather, dim_date, dim_cancellation_reason) + Data Mart mart_delay_analysis.")

    # --- SECTION 4 ---
    add_heading_1("4. TỔNG QUAN MÃ NGUỒN THƯ MỤC `src/` (TRÁI TIM ĐIỀU PHỐI ELT PIPELINE)")
    add_body("Toàn bộ quy trình cào nạp, xử lý và kiểm định dữ liệu của AeroFlow được đóng gói modular chuyên nghiệp trong thư mục `src/`. Mỗi file Python giữ một vai trò nhiệm vụ ELT riêng biệt:")

    tbl_src = doc.add_table(rows=7, cols=3)
    tbl_src.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_src = ["Tên File trong `src/`", "Giai Đoạn ELT Đảm Nhận", "Chi Tiết Nhiệm Vụ & Thuật Toán"]
    hdr_cells_src = tbl_src.rows[0].cells
    for idx, header_text in enumerate(headers_src):
        hdr_cells_src[idx].text = header_text
        set_cell_background(hdr_cells_src[idx], "0A2540")
        p = hdr_cells_src[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)

    src_rows = [
        ("src/config.py", "Configuration", "Cấu hình trung tâm: Lưu tham số GCP Project ID, Bucket Data Lake, BQ Dataset (staging/warehouse), Endpoints API & URL dữ liệu tĩnh."),
        ("src/extract.py", "Phase E (Extract)", "Trích xuất dữ liệu đa nguồn: Cào AviationStack API với thuật toán Auto Dual-Key Failover Rotation, Open-Meteo API, đẩy file thô lên GCS Bucket."),
        ("src/load.py", "Phase L (Load)", "Nạp dữ liệu thô: Đọc file thô từ GCS Data Lake, cấu hình Load Job đẩy vào BigQuery Staging ở chế độ WRITE_APPEND (nạp nối tiếp)."),
        ("src/transform.py", "Phase T (Transform)", "Biến đổi dữ liệu chuyên sâu 9 bước SQL: Khử trùng lặp Staging, Dim Airport, Dim Date, Dim Weather, Dim Carrier, Dim Aircraft (FARM_FINGERPRINT), Fact Flights (SHA256, Partition, Cluster), Mart Delay Analysis."),
        ("src/quality.py", "Phase Q (Quality)", "Kiểm định chất lượng 5 bài test: PK Unique/Non-Null Audit, FK Referential Integrity Audit, Domain Validity Audit (distance_miles > 0)."),
        ("src/orchestrator.py", "Orchestration & Log", "Điều phối luồng chạy E -> L -> T -> Q và tự động lưu nhật ký audit log vĩnh viễn (Append-only) vào BigQuery staging.etl_execution_logs.")
    ]

    for row_idx, data in enumerate(src_rows, 1):
        row_cells = tbl_src.rows[row_idx].cells
        bg_color = "F9FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row_cells[col_idx]
            cell.text = text
            set_cell_background(cell, bg_color)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.size = Pt(8.5)
                r.font.color.rgb = DARK_TEXT

    # --- SECTION 5 ---
    add_heading_1("5. PHA EXTRACTION (E - `src/extract.py`) - WORKFLOW & DUAL-KEY FAILOVER")
    add_body("File `src/extract.py` cào nạp dữ liệu chuyến bay và thời tiết thời gian thực. Tích hợp thuật toán Xoay Key tự động (Dual-Key Rotation: Key 1 -> Key 2 215a8f99667aacc60b9ec21fa615e363) giúp cào liên tục 100% không sợ hết Quota.")
    add_code_block("""[EXTRACTION WORKFLOW]
AviationStack API (Key 1) ──► (Nếu Lỗi/Hết Quota) ──► Chuyển Key 2 Dự Phòng ──► GCS Data Lake (raw/)""")

    # --- SECTION 6 ---
    add_heading_1("6. PHA LOADING (L - `src/load.py`) - NẠP NỐI TIẾP VÀO STAGING")
    add_body("File `src/load.py` thực thi BigQuery Load Job đẩy dữ liệu từ GCS vào BigQuery Staging ở chế độ WRITE_APPEND (Nạp nối tiếp cộng dồn).")
    add_code_block("""[LOADING WORKFLOW]
GCS Data Lake (raw/) ──► BigQuery Load Job Config (WRITE_APPEND) ──► staging.stg_flights_raw""")

    # --- SECTION 7 ---
    add_heading_1("7. PHA TRANSFORMATION (T - `src/transform.py`) - 9 BƯỚC SQL TRONG BIGQUERY")
    add_body("File `src/transform.py` thực thi 9 bước SQL biến đổi dữ liệu trực tiếp trong kho BigQuery:")
    add_bullet("Bước 0 (deduplicate_staging_flights): ROW_NUMBER() OVER (PARTITION BY FlightDate, Airline, FlightNum, Origin, Dest ORDER BY DepTime IS NOT NULL DESC).", "• ")
    add_bullet("Bước 1 (transform_dim_airport): FULL OUTER JOIN giữa OurAirports & OpenFlights, LEFT JOIN đếm số đường băng và MAX(length_ft).", "• ")
    add_bullet("Bước 2 (transform_dim_date): UNNEST(GENERATE_DATE_ARRAY('2026-01-01', '2026-12-31')), phân loại Ngày cuối tuần & Mùa.", "• ")
    add_bullet("Bước 3 (transform_dim_weather): Phân loại điều kiện thời tiết (Heavy Rain, Rain, Snow, Clear) & QUALIFY ROW_NUMBER() = 1.", "• ")
    add_bullet("Bước 4 (transform_dim_carrier): Phân nhóm Hãng lớn (Major Legacy), Hãng giá rẻ (Low-Cost), Hãng khu vực (Regional).", "• ")
    add_bullet("Bước 5 (transform_dim_aircraft): Thuật toán băm FARM_FINGERPRINT phân loại Boeing vs Airbus & tính toán tuổi tàu bay.", "• ")
    add_bullet("Bước 6 (transform_dim_cancellation_reason): Danh mục 5 mã hủy chuyến chuẩn FAA/BTS Mỹ bằng UNNEST STRUCT.", "• ")
    add_bullet("Bước 7 (transform_fact_flights): Mã hóa SHA256(Tail_Number) bảo mật. PARTITION BY flight_date & CLUSTER BY carrier_key, origin_airport_key, dest_airport_key (giảm 95% chi phí truy vấn).", "• ")
    add_bullet("Bước 8 (transform_mart_delay_analysis): Bảng Data Mart phục vụ báo cáo Dashboard siêu tốc.", "• ")

    # --- SECTION 8 ---
    add_heading_1("8. PHA QUALITY AUDIT (Q - `src/quality.py`) & HỆ THỐNG LOGGING (`src/orchestrator.py`)")
    add_body("File `src/quality.py` kiểm định 5 bài test (PASS 100%): PK Unique/Non-Null, FK Referential Integrity, Domain Validity. File `src/orchestrator.py` điều phối luồng E -> L -> T -> Q và ghi nhật ký nối tiếp vĩnh viễn vào BigQuery staging.etl_execution_logs bằng bigquery.ScalarQueryParameter phân biệt ⚡ Nạp Thủ Công Admin và 🤖 Tự Động Cron Job.")

    # --- SECTION 9 ---
    add_heading_1("9. DEMO DATA PIPELINE & ỨNG DỤNG WEB PORTAL 2 CỔNG")
    add_body("Ứng dụng Web Portal triển khai trên Cloud Run phục vụ Cổng Khách Hàng (Thuật toán sắp xếp vòng lặp 24h MOD(CRS_DepTime - CurrentTime + 2400, 2400) ASC, Risk Score) và Cổng Quản Trị Viên (Nút ⚡ Cào Nạp Daily, Modal Quality Audit 5 bài test, Panel Minh bạch Lineage, Biểu đồ Boeing vs Airbus và Tự động làm mới ngầm 15 giây).")

    # Document Footer
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("--- HẾT BÁO CÁO THUYẾT MINH THIẾT KẾ ĐỒ ÁN AEROFLOW GCP PIPELINE ---")
    r_foot.font.italic = True
    r_foot.font.bold = True
    r_foot.font.color.rgb = SECONDARY_COLOR

    doc_path = r"d:\HocTap\CAP2\BaoCao_DoAn_AeroFlow_GCP.docx"
    doc.save(doc_path)
    print(f"✔ Đã tạo thành công file Word Báo cáo Thuyết minh Thiết kế Đồ án tại: {doc_path}")

if __name__ == "__main__":
    create_report()
